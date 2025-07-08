from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from string import ascii_lowercase
import pandas as pd

def generate_suffixes(n):
    suffixes = []
    reps = 1
    while len(suffixes) < n:
        for ch in ascii_lowercase:
            suffixes.append(ch * reps)
            if len(suffixes) == n:
                break
        reps += 1
    return suffixes

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Tahoma")
    rPr.append(font)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rPr.append(underline)

    run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def generate_reference_docx(df, prefix, agency_name, proceeding_code):
    df = df[['TN #', 'Docketed Date', 'Document Title']].dropna()
    df['Year'] = pd.to_datetime(df['Docketed Date'], errors='coerce').dt.year
    df['Formatted Title'] = df['Document Title'].apply(lambda x: str(x).split('\n')[0].strip())
    df = df.sort_values(by=['Year', 'Docketed Date']).reset_index(drop=True)
    df['GroupIndex'] = df.groupby('Year').cumcount()
    suffix_dict = {y: generate_suffixes(len(g)) for y, g in df.groupby('Year')}
    df['Suffix'] = df.apply(lambda r: suffix_dict[r['Year']][r['GroupIndex']], axis=1)
    df['Formatted Date'] = pd.to_datetime(df['Docketed Date'], errors='coerce').dt.strftime('%B %#d, %Y')

    base_url = f"https://efiling.energy.ca.gov/Lists/DocketLog.aspx?docketnumber={proceeding_code}"

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Tahoma'
    style.font.size = Pt(12)
    para_format = style.paragraph_format
    para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)

    for _, row in df.iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"{prefix} {row['Year']}{row['Suffix']} – {agency_name} (TN {row['TN #']}). {row['Formatted Title']}. Docketed {row['Formatted Date']}. Accessed online at: ")
        run.font.name = 'Tahoma'
        run.font.size = Pt(12)
        run.font.underline = False
        run.font.color.rgb = None
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        add_hyperlink(p, base_url, base_url)

    output_path = "Reference_List.docx"
    doc.save(output_path)
    return output_path
